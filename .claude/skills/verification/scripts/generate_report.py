#!/usr/bin/env python3
"""
generate_report.py — Генерация итогового отчёта верификации.
Принимает результаты скриптов (JSON) и LLM-проверок (md), сводит в единый отчёт.

Использование:
    python generate_report.py --script-results script_report.json --output report.md
    python generate_report.py --script-results script_report.json --agent-results agent.md --output report.md
    python generate_report.py --script-results script_report.json --output report.docx --format docx
"""

import argparse, json, re, sys
from datetime import datetime
from pathlib import Path


def load_script_results(filepath):
    if not filepath or not Path(filepath).exists():
        return None
    try:
        return json.loads(Path(filepath).read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        print(f"⚠️ Ошибка разбора JSON: {filepath} — {e}", file=sys.stderr)
        return None


def load_agent_results(filepath):
    if not filepath or not Path(filepath).exists():
        return None
    return Path(filepath).read_text(encoding="utf-8")


def generate_md(script_results, agent_results, preset=None, mechanisms=None):
    lines = [
        "# 📊 Отчёт верификации",
        "",
        f"**Дата:** {datetime.now().strftime('%d.%m.%Y %H:%M')}",
    ]

    if preset:
        lines.append(f"**Пресет:** {preset}")
    if mechanisms:
        lines.append(f"**Механизмы:** {mechanisms}")
    lines.append("")

    # Скриптовые результаты
    if script_results:
        lines.extend([
            "## Скриптовые проверки (автоматические)",
            "",
            "| Скрипт | Статус | Проверено | ⚠️ | ❌ | Детали |",
            "|--------|--------|-----------|-----|-----|--------|",
        ])

        for check in script_results.get("checks", []):
            status_icon = {"pass": "✅", "warn": "⚠️", "fail": "❌", "skip": "⏭️", "error": "💥", "info": "ℹ️"}.get(
                check.get("status", "?"), "?"
            )
            lines.append(
                f"| {check.get('script', '?')} | {status_icon} | "
                f"{check.get('items_checked', 0)} | {check.get('items_warned', 0)} | "
                f"{check.get('items_failed', 0)} | {check.get('details', '')} |"
            )

        # Сводка
        summary = script_results.get("summary", {})
        lines.extend([
            "",
            f"**Итого скрипты:** проверено {summary.get('total_checked', 0)} | "
            f"✅ {summary.get('total_passed', 0)} | ⚠️ {summary.get('total_warned', 0)} | "
            f"❌ {summary.get('total_failed', 0)}",
            "",
        ])

        # Детали замечаний
        all_findings = []
        for check in script_results.get("checks", []):
            all_findings.extend(check.get("findings", []))

        if all_findings:
            lines.extend(["### Детали замечаний", ""])
            for f in all_findings:
                if f.get("severity") in ("warning", "error"):
                    sev = {"warning": "⚠️", "error": "❌"}.get(f["severity"], "?")
                    lines.append(f"- {sev} **{f.get('location', '?')}**: {f.get('description', '?')}")
            lines.append("")

    # LLM-результаты
    if agent_results:
        lines.extend([
            "## LLM-проверки (агентные)",
            "",
            agent_results,
            "",
        ])

    # Общий вердикт
    overall = "pass"
    if script_results:
        s = script_results.get("summary", {}).get("overall_status", "pass")
        if s in ("fail", "warn"):
            overall = s

    verdict = {"pass": "✅ Верификация пройдена", "warn": "⚠️ Верификация пройдена с замечаниями",
               "fail": "❌ Верификация выявила ошибки"}.get(overall, "?")
    
    lines.extend([
        "---",
        f"## Итоговый вердикт: {verdict}",
        "",
    ])

    return "\n".join(lines)


def _add_formatted_paragraph(doc, text, style="Normal"):
    """Добавить абзац с парсингом inline-разметки (**жирный**, *курсив*)."""
    p = doc.add_paragraph(style=style)
    # Разбиваем по **жирный** и *курсив*
    # Порядок: сначала bold+italic (***), потом bold (**), потом italic (*)
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)')
    last_end = 0
    for m in pattern.finditer(text):
        # Добавляем текст до матча
        if m.start() > last_end:
            p.add_run(text[last_end:m.start()])
        if m.group(2):  # ***bold+italic***
            run = p.add_run(m.group(2))
            run.bold = True
            run.italic = True
        elif m.group(3):  # **bold**
            run = p.add_run(m.group(3))
            run.bold = True
        elif m.group(4):  # *italic*
            run = p.add_run(m.group(4))
            run.italic = True
        last_end = m.end()
    # Остаток текста
    if last_end < len(text):
        p.add_run(text[last_end:])
    return p


def generate_docx(md_content, output_path):
    """Конвертировать md-отчёт в docx с нормальными таблицами."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(1.5)

        lines = md_content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("---"):
                doc.add_paragraph("─" * 60)
            elif line.startswith("|"):
                # Собираем все строки таблицы
                table_lines = []
                while i < len(lines) and lines[i].startswith("|"):
                    row_text = lines[i]
                    # Пропускаем строки-разделители (|---|---|)
                    if not set(row_text.replace("|", "").strip()) <= {"-", " ", ":"}:
                        # Разделяем по |, но учитываем экранированный \|
                        raw = row_text.strip().strip("|")
                        cells = [c.strip().replace("\\|", "|") for c in re.split(r'(?<!\\)\|', raw)]
                        table_lines.append(cells)
                    i += 1
                # Создаём docx-таблицу
                if table_lines:
                    num_cols = max(len(row) for row in table_lines)
                    tbl = doc.add_table(rows=len(table_lines), cols=num_cols, style="Table Grid")
                    for r_idx, row_cells in enumerate(table_lines):
                        for c_idx, cell_text in enumerate(row_cells):
                            if c_idx < num_cols:
                                cell = tbl.cell(r_idx, c_idx)
                                cell.text = cell_text
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.font.size = Pt(10)
                    # Жирный шрифт для заголовков (первая строка)
                    if table_lines:
                        for c_idx in range(min(len(table_lines[0]), num_cols)):
                            cell = tbl.cell(0, c_idx)
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.bold = True
                i -= 1  # компенсация внешнего i += 1
            elif line.startswith("- "):
                _add_formatted_paragraph(doc, line[2:], style="List Bullet")
            elif line.strip():
                _add_formatted_paragraph(doc, line)

            i += 1

        doc.save(output_path)
        return True
    except ImportError:
        return False


def main():
    parser = argparse.ArgumentParser(description="Генерация отчёта верификации")
    parser.add_argument("--script-results", default=None, help="JSON с результатами скриптов")
    parser.add_argument("--agent-results", default=None, help="MD с результатами LLM-агента")
    parser.add_argument("--preset", default=None, help="Название пресета")
    parser.add_argument("--mechanisms", default=None, help="Список механизмов")
    parser.add_argument("--output", required=True, help="Путь к выходному файлу")
    parser.add_argument("--format", choices=["md", "docx"], default=None, help="Формат (авто по расширению)")

    args = parser.parse_args()

    script_results = load_script_results(args.script_results)
    agent_results = load_agent_results(args.agent_results)

    md_content = generate_md(script_results, agent_results, args.preset, args.mechanisms)

    output_format = args.format or ("docx" if args.output.endswith(".docx") else "md")

    if output_format == "docx":
        if generate_docx(md_content, args.output):
            print(f"💾 Отчёт сохранён: {args.output} (docx)")
        else:
            # Fallback to md
            md_path = args.output.replace(".docx", ".md")
            Path(md_path).write_text(md_content, encoding="utf-8")
            print(f"💾 Отчёт сохранён: {md_path} (md, python-docx не установлен)")
    else:
        Path(args.output).write_text(md_content, encoding="utf-8")
        print(f"💾 Отчёт сохранён: {args.output} (md)")


if __name__ == "__main__":
    main()
