#!/usr/bin/env python3
"""
verify_references.py — Поиск битых внутренних ссылок в docx.
Находит паттерны: «см. п. X», «Таблица N», «Приложение X» и проверяет существование целей.
"""

import argparse, json, re, sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("pip install python-docx --break-system-packages", file=sys.stderr)
    sys.exit(1)

REF_PATTERNS = [
    (r"(?:см\.?\s*)?(?:таблиц[аеуы]\s*)(\d+)", "Таблица"),
    (r"(?:см\.?\s*)?(?:рисун\w*\s*)(\d+)", "Рисунок"),
    (r"(?:см\.?\s*)?(?:диаграмм[аеуы]\s*)(\d+)", "Диаграмма"),
    (r"(?:см\.?\s*)?(?:приложени[еяю]\s*)([А-ЯA-Z\d]+)", "Приложение"),
    (r"(?:см\.?\s*)?(?:п(?:ункт)?\.?\s*)(\d+(?:\.\d+)*)", "Пункт"),
    (r"(?:на\s+)?(?:слайд[еу]?\s*)(\d+)", "Слайд"),
    (r"(?:см\.?\s*)?(?:схем[аеуы]\s*)(\d+)", "Схема"),
    (r"(?:см\.?\s*)?(?:график[аеу]?\s*)(\d+)", "График"),
]

def extract_references(doc):
    """Извлечь все ссылки из текста, обрабатывая поабзацно."""
    refs = []
    for para_idx, para in enumerate(doc.paragraphs, 1):
        text = para.text
        if not text.strip():
            continue
        for pattern, ref_type in REF_PATTERNS:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                context = text[max(0, match.start()-30):match.end()+30].strip()
                refs.append({
                    "type": ref_type,
                    "target": match.group(1),
                    "context": context,
                    "paragraph": para_idx,
                })
    return refs

def find_targets(doc):
    """Найти существующие цели (таблицы, заголовки, приложения)."""
    targets = {"Таблица": set(), "Рисунок": set(), "Диаграмма": set(),
               "Приложение": set(), "Пункт": set(), "Схема": set(), "График": set(), "Слайд": set()}
    
    # Таблицы — считаем по порядку
    for i in range(len(doc.tables)):
        targets["Таблица"].add(str(i + 1))

    # Заголовки с нумерацией → пункты
    for para in doc.paragraphs:
        text = para.text.strip()
        # Пункты (1.1, 2.3.1, etc.)
        m = re.match(r"^(\d+(?:\.\d+)*)[.\s]", text)
        if m:
            targets["Пункт"].add(m.group(1))
        # Приложения
        m = re.match(r"^Приложение\s+([А-ЯA-Z\d]+)", text, re.IGNORECASE)
        if m:
            targets["Приложение"].add(m.group(1).upper())
        # Подписи к рисункам/диаграммам/схемам
        for obj_type in ["Рисунок", "Диаграмма", "Схема", "График"]:
            pattern = rf"^{obj_type}\s+(\d+)"
            m = re.match(pattern, text, re.IGNORECASE)
            if m:
                targets[obj_type].add(m.group(1))

    return targets

def verify(filepath: str) -> dict:
    doc = Document(filepath)
    refs = extract_references(doc)
    targets = find_targets(doc)
    findings = []

    for ref in refs:
        target_set = targets.get(ref["type"], set())
        target_id = ref["target"].upper() if ref["type"] == "Приложение" else ref["target"]
        
        if target_id not in target_set:
            para_info = f" (абзац {ref['paragraph']})" if "paragraph" in ref else ""
            findings.append({
                "severity": "error",
                "location": f"{ref['type']} {ref['target']}{para_info}",
                "expected": "существует",
                "actual": "не найдено",
                "description": f"Битая ссылка: {ref['type']} {ref['target']} не найден(а) в документе. Контекст: «...{ref['context']}...»",
            })

    items_checked = len(refs)
    items_failed = len(findings)
    status = "fail" if items_failed > 0 else "pass"

    return {
        "script": "verify_references.py",
        "status": status,
        "details": f"Ссылок найдено: {items_checked}, битых: {items_failed}",
        "items_checked": items_checked,
        "items_passed": items_checked - items_failed,
        "items_warned": 0,
        "items_failed": items_failed,
        "findings": findings,
    }

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("files", nargs="+")
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()
    for f in args.files:
        if f.endswith(".docx"):
            result = verify(f)
            print(json.dumps(result, ensure_ascii=False, indent=2) if args.json else f"🔗 {f}: {result['status']} ({result['items_failed']} битых)")
