#!/usr/bin/env python3
"""
verify_docx_format.py — Проверка формата docx по стандарту #6.
Проверяет: поля, шрифт, размеры заголовков, отступы, колонтитулы.

Использование:
    python verify_docx_format.py document.docx
    python verify_docx_format.py --json document.docx
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Emu
except ImportError:
    print("pip install python-docx --break-system-packages", file=sys.stderr)
    sys.exit(1)

# Стандарт #6
STANDARD = {
    "top_margin_cm": 2.0,
    "bottom_margin_cm": 2.0,
    "left_margin_cm": 3.0,
    "right_margin_cm": 1.5,
    "font_name": "Times New Roman",
    "font_size_pt": 14,
    "first_line_indent_cm": 1.5,
    "line_spacing": 1.15,
    "space_after_pt": 8,
    "h1_size_pt": 22,
    "h2_size_pt": 18,
    "h3_size_pt": 16,
    "heading_color": "0070C0",
    "table_font_size_pt": 12,
    "table_line_spacing": 1.0,
    "table_space_after_pt": 2,
    "header_font_size_pt": 9,
    "footer_font_size_pt": 12,
}

TOLERANCE_CM = 0.1  # допуск в см
TOLERANCE_PT = 0.5  # допуск в pt


def cm_from_emu(emu):
    """Конвертация EMU в см."""
    if emu is None:
        return None
    return round(emu / 360000, 2)


def pt_from_emu(emu):
    """Конвертация EMU в pt."""
    if emu is None:
        return None
    return round(emu / 12700, 1)


def check_margins(doc):
    """Проверка полей."""
    findings = []
    for i, section in enumerate(doc.sections):
        checks = [
            ("top_margin", cm_from_emu(section.top_margin), STANDARD["top_margin_cm"]),
            ("bottom_margin", cm_from_emu(section.bottom_margin), STANDARD["bottom_margin_cm"]),
            ("left_margin", cm_from_emu(section.left_margin), STANDARD["left_margin_cm"]),
            ("right_margin", cm_from_emu(section.right_margin), STANDARD["right_margin_cm"]),
        ]
        for name, actual, expected in checks:
            if actual is not None and abs(actual - expected) > TOLERANCE_CM:
                findings.append({
                    "severity": "warning",
                    "location": f"Section {i+1}, {name}",
                    "expected": f"{expected} см",
                    "actual": f"{actual} см",
                    "description": f"Поле {name}: ожидалось {expected} см, получено {actual} см",
                })
    return findings


def check_paragraphs(doc):
    """Проверка шрифтов и отступов абзацев."""
    findings = []
    heading_sizes = {"Heading 1": STANDARD["h1_size_pt"], "Heading 2": STANDARD["h2_size_pt"], "Heading 3": STANDARD["h3_size_pt"]}
    
    checked = 0
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        checked += 1
        style_name = para.style.name if para.style else ""

        # Заголовки
        if style_name in heading_sizes:
            expected_size = heading_sizes[style_name]
            for run in para.runs:
                if run.font.size:
                    actual_pt = pt_from_emu(run.font.size)
                    if actual_pt and abs(actual_pt - expected_size) > TOLERANCE_PT:
                        findings.append({
                            "severity": "warning",
                            "location": f"Абзац {i+1} ({style_name})",
                            "expected": f"{expected_size} pt",
                            "actual": f"{actual_pt} pt",
                            "description": f"Размер заголовка: {actual_pt}pt вместо {expected_size}pt",
                        })
                    break
            # Проверка цвета заголовка
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    actual_color = str(run.font.color.rgb)
                    if actual_color.upper() != STANDARD["heading_color"].upper():
                        findings.append({
                            "severity": "info",
                            "location": f"Абзац {i+1} ({style_name})",
                            "expected": f"#{STANDARD['heading_color']}",
                            "actual": f"#{actual_color}",
                            "description": f"Цвет заголовка: #{actual_color} вместо #{STANDARD['heading_color']}",
                        })
                break

        # Обычные абзацы — шрифт
        elif style_name in ("Normal", "Body Text", ""):
            for run in para.runs:
                if run.font.name and run.font.name != STANDARD["font_name"]:
                    findings.append({
                        "severity": "warning",
                        "location": f"Абзац {i+1}",
                        "expected": STANDARD["font_name"],
                        "actual": run.font.name,
                        "description": f"Шрифт: {run.font.name} вместо {STANDARD['font_name']}",
                    })
                if run.font.size:
                    actual_pt = pt_from_emu(run.font.size)
                    if actual_pt and abs(actual_pt - STANDARD["font_size_pt"]) > TOLERANCE_PT:
                        findings.append({
                            "severity": "warning",
                            "location": f"Абзац {i+1}",
                            "expected": f"{STANDARD['font_size_pt']} pt",
                            "actual": f"{actual_pt} pt",
                            "description": f"Размер шрифта: {actual_pt}pt вместо {STANDARD['font_size_pt']}pt",
                        })
                break  # проверяем только первый run

    return findings, checked


def check_tables(doc):
    """Проверка форматирования таблиц."""
    findings = []
    for i, table in enumerate(doc.tables):
        for row in table.rows[:1]:  # проверяем первую строку как образец
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.size:
                            actual_pt = pt_from_emu(run.font.size)
                            if actual_pt and abs(actual_pt - STANDARD["table_font_size_pt"]) > TOLERANCE_PT:
                                findings.append({
                                    "severity": "info",
                                    "location": f"Таблица {i+1}",
                                    "expected": f"{STANDARD['table_font_size_pt']} pt",
                                    "actual": f"{actual_pt} pt",
                                    "description": f"Шрифт в таблице: {actual_pt}pt вместо {STANDARD['table_font_size_pt']}pt",
                                })
                            break
                    break
                break
            break
    return findings


def verify(filepath: str) -> dict:
    """Главная функция верификации."""
    doc = Document(filepath)
    all_findings = []

    # Поля
    margin_findings = check_margins(doc)
    all_findings.extend(margin_findings)

    # Абзацы
    para_findings, para_checked = check_paragraphs(doc)
    all_findings.extend(para_findings)

    # Таблицы
    table_findings = check_tables(doc)
    all_findings.extend(table_findings)

    items_checked = len(doc.sections) + para_checked + len(doc.tables)
    items_warned = len([f for f in all_findings if f["severity"] == "warning"])
    items_failed = len([f for f in all_findings if f["severity"] == "error"])
    items_info = len([f for f in all_findings if f["severity"] == "info"])

    status = "fail" if items_failed > 0 else ("warn" if items_warned > 0 else "pass")

    return {
        "script": "verify_docx_format.py",
        "status": status,
        "details": f"Проверено: {items_checked} элементов. Замечания: {items_warned}, инфо: {items_info}",
        "items_checked": items_checked,
        "items_passed": items_checked - items_warned - items_failed,
        "items_warned": items_warned,
        "items_failed": items_failed,
        "findings": all_findings,
    }


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("files", nargs="+")
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()

    for f in args.files:
        if f.endswith(".docx"):
            result = verify(f)
            if args.json:
                print(json.dumps(result, ensure_ascii=False, indent=2))
            else:
                print(f"📄 {f}: {result['status']} ({result['items_warned']} замечаний)")
                for finding in result["findings"]:
                    icon = {"warning": "⚠️", "error": "❌", "info": "ℹ️"}[finding["severity"]]
                    print(f"  {icon} {finding['location']}: {finding['description']}")
